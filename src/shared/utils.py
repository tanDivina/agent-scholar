"""
Shared utility functions for Agent Scholar

This module contains common utility functions used across
all Lambda functions and components.
"""

import json
import logging
import hashlib
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
import boto3
from botocore.exceptions import ClientError

# Configure logging
logger = logging.getLogger(__name__)

def setup_logging(level: str = "INFO") -> None:
    """
    Set up logging configuration.
    
    Args:
        level: Logging level (DEBUG, INFO, WARNING, ERROR)
    """
    logging.basicConfig(
        level=getattr(logging, level.upper()),
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

def generate_id(prefix: str = "") -> str:
    """
    Generate a unique identifier.
    
    Args:
        prefix: Optional prefix for the ID
        
    Returns:
        Unique identifier string
    """
    unique_id = str(uuid.uuid4())
    return f"{prefix}{unique_id}" if prefix else unique_id

def hash_content(content: str) -> str:
    """
    Generate a hash for content deduplication.
    
    Args:
        content: Content string to hash
        
    Returns:
        SHA-256 hash of the content
    """
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200, 
               preserve_sentences: bool = True, preserve_paragraphs: bool = True) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks for embedding with intelligent boundary detection.
    
    Args:
        text: Text to chunk
        chunk_size: Maximum size of each chunk in characters
        overlap: Number of characters to overlap between chunks
        preserve_sentences: Whether to try to break at sentence boundaries
        preserve_paragraphs: Whether to try to break at paragraph boundaries
        
    Returns:
        List of chunk dictionaries with content and position info
    """
    if not text or not text.strip():
        return []
    
    text = text.strip()
    
    if len(text) <= chunk_size:
        return [{
            'content': text,
            'start_position': 0,
            'end_position': len(text),
            'chunk_index': 0,
            'word_count': len(text.split()),
            'char_count': len(text)
        }]
    
    chunks = []
    start = 0
    chunk_index = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        
        # If we're not at the end of the text, try to find a good break point
        if end < len(text):
            end = _find_optimal_break_point(text, start, end, preserve_sentences, preserve_paragraphs)
        
        chunk_content = text[start:end].strip()
        
        if chunk_content:
            chunks.append({
                'content': chunk_content,
                'start_position': start,
                'end_position': end,
                'chunk_index': chunk_index,
                'word_count': len(chunk_content.split()),
                'char_count': len(chunk_content)
            })
            chunk_index += 1
        
        # Calculate next start position with overlap
        if end >= len(text):
            break
            
        # Find overlap start position
        overlap_start = max(start + 1, end - overlap)
        
        # Try to start overlap at a word boundary
        if overlap_start < end:
            word_boundary = text.rfind(' ', overlap_start, end)
            if word_boundary > overlap_start:
                overlap_start = word_boundary + 1
        
        start = overlap_start
    
    return chunks


def _find_optimal_break_point(text: str, start: int, max_end: int, 
                             preserve_sentences: bool, preserve_paragraphs: bool) -> int:
    """
    Find the optimal break point for text chunking.
    
    Args:
        text: Full text
        start: Start position of current chunk
        max_end: Maximum end position
        preserve_sentences: Whether to preserve sentence boundaries
        preserve_paragraphs: Whether to preserve paragraph boundaries
        
    Returns:
        Optimal end position for the chunk
    """
    # Define the search window (last 20% of the chunk)
    search_start = max(start, max_end - (max_end - start) // 5)
    
    # First priority: paragraph boundaries
    if preserve_paragraphs:
        paragraph_end = text.rfind('\n\n', search_start, max_end)
        if paragraph_end > search_start:
            return paragraph_end + 2
        
        # Also look for single newlines
        newline_end = text.rfind('\n', search_start, max_end)
        if newline_end > search_start:
            return newline_end + 1
    
    # Second priority: sentence boundaries
    if preserve_sentences:
        # Look for sentence endings
        sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
        best_sentence_end = -1
        
        for ending in sentence_endings:
            pos = text.rfind(ending, search_start, max_end)
            if pos > best_sentence_end:
                best_sentence_end = pos + len(ending)
        
        if best_sentence_end > search_start:
            return best_sentence_end
    
    # Third priority: word boundaries
    word_boundary = text.rfind(' ', search_start, max_end)
    if word_boundary > search_start:
        return word_boundary
    
    # Fallback: use max_end
    return max_end


def chunk_text_semantic(text: str, max_chunk_size: int = 1000, min_chunk_size: int = 100,
                       overlap_ratio: float = 0.1) -> List[Dict[str, Any]]:
    """
    Advanced semantic chunking that tries to keep related content together.
    
    Args:
        text: Text to chunk
        max_chunk_size: Maximum chunk size in characters
        min_chunk_size: Minimum chunk size in characters
        overlap_ratio: Ratio of overlap between chunks (0.0 to 0.5)
        
    Returns:
        List of chunk dictionaries with semantic boundaries
    """
    if not text or not text.strip():
        return []
    
    text = text.strip()
    
    # First, split by paragraphs
    paragraphs = text.split('\n\n')
    
    chunks = []
    current_chunk = ""
    current_start = 0
    chunk_index = 0
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If adding this paragraph would exceed max size, finalize current chunk
        if current_chunk and len(current_chunk) + len(paragraph) + 2 > max_chunk_size:
            if len(current_chunk) >= min_chunk_size:
                chunks.append({
                    'content': current_chunk.strip(),
                    'start_position': current_start,
                    'end_position': current_start + len(current_chunk),
                    'chunk_index': chunk_index,
                    'word_count': len(current_chunk.split()),
                    'char_count': len(current_chunk),
                    'semantic_type': 'paragraph_boundary'
                })
                chunk_index += 1
                
                # Calculate overlap
                overlap_size = int(len(current_chunk) * overlap_ratio)
                if overlap_size > 0:
                    overlap_text = current_chunk[-overlap_size:].strip()
                    current_chunk = overlap_text + '\n\n' + paragraph
                else:
                    current_chunk = paragraph
                    
                current_start = current_start + len(current_chunk) - len(paragraph) - 2
            else:
                current_chunk += '\n\n' + paragraph
        else:
            if current_chunk:
                current_chunk += '\n\n' + paragraph
            else:
                current_chunk = paragraph
    
    # Add the last chunk
    if current_chunk and len(current_chunk.strip()) >= min_chunk_size:
        chunks.append({
            'content': current_chunk.strip(),
            'start_position': current_start,
            'end_position': current_start + len(current_chunk),
            'chunk_index': chunk_index,
            'word_count': len(current_chunk.split()),
            'char_count': len(current_chunk),
            'semantic_type': 'final_chunk'
        })
    
    # If no chunks were created (text too short), create one chunk
    if not chunks and text:
        chunks.append({
            'content': text,
            'start_position': 0,
            'end_position': len(text),
            'chunk_index': 0,
            'word_count': len(text.split()),
            'char_count': len(text),
            'semantic_type': 'single_chunk'
        })
    
    return chunks

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text content from various file formats.
    
    Args:
        file_path: Path to the file
        file_type: Type of file (pdf, docx, txt, html, md)
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
        Exception: For other extraction errors
    """
    import os
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_type = file_type.lower().strip('.')
    
    try:
        if file_type == 'txt':
            return _extract_text_from_txt(file_path)
        elif file_type == 'pdf':
            return _extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return _extract_text_from_docx(file_path)
        elif file_type == 'html':
            return _extract_text_from_html(file_path)
        elif file_type in ['md', 'markdown']:
            return _extract_text_from_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        raise


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    # Clean up the text
                    return _clean_extracted_text(content)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            content = f.read().decode('utf-8', errors='ignore')
            return _clean_extracted_text(content)
            
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {str(e)}")
        raise


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file.
    
    Note: This is a basic implementation. In production, you might want to use
    libraries like PyPDF2, pdfplumber, or pymupdf for better PDF handling.
    """
    try:
        # For now, we'll use a simple approach that works in Lambda environment
        # In production, you'd install and use proper PDF libraries
        
        # Try to import PyPDF2 if available
        try:
            import PyPDF2
            return _extract_with_pypdf2(file_path)
        except ImportError:
            pass
        
        # Try to import pdfplumber if available
        try:
            import pdfplumber
            return _extract_with_pdfplumber(file_path)
        except ImportError:
            pass
        
        # Fallback: return a message indicating PDF processing is not available
        logger.warning(f"PDF processing libraries not available for {file_path}")
        return f"[PDF content from {file_path} - PDF processing libraries not installed]"
        
    except Exception as e:
        logger.error(f"Error extracting PDF {file_path}: {str(e)}")
        raise


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Note: This requires python-docx library.
    """
    try:
        # Try to import python-docx if available
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            content = '\n'.join(text_content)
            return _clean_extracted_text(content)
            
        except ImportError:
            logger.warning(f"python-docx library not available for {file_path}")
            return f"[DOCX content from {file_path} - python-docx library not installed]"
            
    except Exception as e:
        logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
        raise


def _extract_text_from_html(file_path: str) -> str:
    """Extract text from HTML file."""
    try:
        # Try to import BeautifulSoup if available
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            return _clean_extracted_text(text)
            
        except ImportError:
            # Fallback: simple HTML tag removal
            import re
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', html_content)
            # Decode HTML entities
            import html
            text = html.unescape(text)
            
            return _clean_extracted_text(text)
            
    except Exception as e:
        logger.error(f"Error extracting HTML {file_path}: {str(e)}")
        raise


def _extract_text_from_markdown(file_path: str) -> str:
    """Extract text from Markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Try to import markdown if available for better processing
        try:
            import markdown
            from markdown.extensions import codehilite
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(content, extensions=['codehilite'])
            
            # Remove HTML tags
            import re
            text = re.sub(r'<[^>]+>', '', html)
            
            return _clean_extracted_text(text)
            
        except ImportError:
            # Fallback: basic markdown processing
            import re
            
            # Remove markdown syntax
            text = re.sub(r'#{1,6}\s+', '', content)  # Headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Inline code
            text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Code blocks
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
            
            return _clean_extracted_text(text)
            
    except Exception as e:
        logger.error(f"Error extracting Markdown {file_path}: {str(e)}")
        raise


def _extract_with_pypdf2(file_path: str) -> str:
    """Extract text using PyPDF2."""
    import PyPDF2
    
    text_content = []
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_content.append(page.extract_text())
    
    content = '\n'.join(text_content)
    return _clean_extracted_text(content)


def _extract_with_pdfplumber(file_path: str) -> str:
    """Extract text using pdfplumber."""
    import pdfplumber
    
    text_content = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
    
    content = '\n'.join(text_content)
    return _clean_extracted_text(content)


def _clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text."""
    import re
    
    if not text:
        return ""
    
    # Remove control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    # Normalize excessive whitespace within lines (but preserve single spaces)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove trailing whitespace from each line
    text = re.sub(r'[ \t]+\n', '\n', text)
    
    # Remove leading whitespace from each line (except first line)
    text = re.sub(r'\n[ \t]+', '\n', text)
    
    # Remove excessive newlines (3 or more consecutive newlines become 2)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace from entire text
    text = text.strip()
    
    return text

def format_timestamp(dt: datetime) -> str:
    """
    Format datetime for consistent display.
    
    Args:
        dt: Datetime object to format
        
    Returns:
        Formatted timestamp string
    """
    return dt.strftime("%Y-%m-%d %H:%M:%S UTC")

def safe_json_loads(json_str: str, default: Any = None) -> Any:
    """
    Safely parse JSON string with fallback.
    
    Args:
        json_str: JSON string to parse
        default: Default value if parsing fails
        
    Returns:
        Parsed JSON object or default value
    """
    try:
        return json.loads(json_str)
    except (json.JSONDecodeError, TypeError) as e:
        logger.warning(f"Failed to parse JSON: {str(e)}")
        return default

def safe_json_dumps(obj: Any, default: str = "{}") -> str:
    """
    Safely serialize object to JSON with fallback.
    
    Args:
        obj: Object to serialize
        default: Default JSON string if serialization fails
        
    Returns:
        JSON string or default value
    """
    try:
        return json.dumps(obj, default=str, ensure_ascii=False)
    except (TypeError, ValueError) as e:
        logger.warning(f"Failed to serialize to JSON: {str(e)}")
        return default

def get_aws_client(service_name: str, region: str = None) -> Any:
    """
    Get AWS service client with error handling.
    
    Args:
        service_name: AWS service name (e.g., 'bedrock', 's3')
        region: AWS region (optional)
        
    Returns:
        AWS service client
    """
    try:
        if region:
            return boto3.client(service_name, region_name=region)
        else:
            return boto3.client(service_name)
    except Exception as e:
        logger.error(f"Failed to create AWS client for {service_name}: {str(e)}")
        raise

def handle_aws_error(error: ClientError, operation: str) -> Dict[str, Any]:
    """
    Handle AWS service errors consistently.
    
    Args:
        error: ClientError from AWS SDK
        operation: Description of the operation that failed
        
    Returns:
        Error response dictionary
    """
    error_code = error.response['Error']['Code']
    error_message = error.response['Error']['Message']
    
    logger.error(f"AWS error in {operation}: {error_code} - {error_message}")
    
    return {
        'error': True,
        'error_code': error_code,
        'error_message': error_message,
        'operation': operation
    }

def validate_environment_variables(required_vars: List[str]) -> bool:
    """
    Validate that required environment variables are set.
    
    Args:
        required_vars: List of required environment variable names
        
    Returns:
        True if all variables are set, False otherwise
    """
    import os
    
    missing_vars = []
    for var in required_vars:
        if not os.getenv(var):
            missing_vars.append(var)
    
    if missing_vars:
        logger.error(f"Missing required environment variables: {missing_vars}")
        return False
    
    return True

def create_bedrock_response(response_text: str, error: bool = False) -> Dict[str, Any]:
    """
    Create standardized response for Bedrock Agent action groups.
    
    Args:
        response_text: Response text to return
        error: Whether this is an error response
        
    Returns:
        Formatted Bedrock Agent response
    """
    return {
        'response': {
            'actionResponse': {
                'actionResponseBody': {
                    'TEXT': {
                        'body': response_text
                    }
                }
            }
        }
    }

def measure_execution_time(func):
    """
    Decorator to measure function execution time.
    
    Args:
        func: Function to measure
        
    Returns:
        Decorated function that logs execution time
    """
    import time
    from functools import wraps
    
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = func(*args, **kwargs)
            execution_time = time.time() - start_time
            logger.info(f"{func.__name__} executed in {execution_time:.3f} seconds")
            return result
        except Exception as e:
            execution_time = time.time() - start_time
            logger.error(f"{func.__name__} failed after {execution_time:.3f} seconds: {str(e)}")
            raise
    
    return wrapper


# Embedding generation functions
def generate_embeddings(texts: List[str], model_name: str = "amazon.titan-embed-text-v1") -> List[List[float]]:
    """
    Generate embeddings for a list of texts using Amazon Bedrock.
    
    Args:
        texts: List of text strings to embed
        model_name: Name of the embedding model to use
        
    Returns:
        List of embedding vectors
        
    Raises:
        Exception: If embedding generation fails
    """
    try:
        bedrock_runtime = get_aws_client('bedrock-runtime')
        embeddings = []
        
        for text in texts:
            if not text or not text.strip():
                # Return zero vector for empty text
                embeddings.append([0.0] * 1536)  # Titan embedding dimension
                continue
                
            embedding = _generate_single_embedding(bedrock_runtime, text.strip(), model_name)
            embeddings.append(embedding)
        
        return embeddings
        
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise


def _generate_single_embedding(bedrock_runtime, text: str, model_name: str) -> List[float]:
    """Generate embedding for a single text."""
    import json
    
    # Prepare the request body
    body = {
        "inputText": text
    }
    
    try:
        response = bedrock_runtime.invoke_model(
            modelId=model_name,
            body=json.dumps(body),
            contentType='application/json',
            accept='application/json'
        )
        
        response_body = json.loads(response['body'].read())
        
        # Extract embedding from response
        if 'embedding' in response_body:
            return response_body['embedding']
        else:
            logger.error(f"No embedding found in response: {response_body}")
            raise ValueError("No embedding found in response")
            
    except Exception as e:
        logger.error(f"Error calling Bedrock embedding model: {str(e)}")
        raise


def generate_embedding_batch(texts: List[str], batch_size: int = 25, 
                           model_name: str = "amazon.titan-embed-text-v1") -> List[List[float]]:
    """
    Generate embeddings in batches to handle rate limits.
    
    Args:
        texts: List of text strings to embed
        batch_size: Number of texts to process in each batch
        model_name: Name of the embedding model to use
        
    Returns:
        List of embedding vectors
    """
    import time
    
    all_embeddings = []
    
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        
        try:
            batch_embeddings = generate_embeddings(batch, model_name)
            all_embeddings.extend(batch_embeddings)
            
            # Add small delay between batches to respect rate limits
            if i + batch_size < len(texts):
                time.sleep(0.1)
                
        except Exception as e:
            logger.error(f"Error processing batch {i//batch_size + 1}: {str(e)}")
            # Add zero vectors for failed batch
            for _ in batch:
                all_embeddings.append([0.0] * 1536)
    
    return all_embeddings


# Document processing pipeline functions
def process_document_for_embedding(file_path: str, document_id: str, 
                                 chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
    """
    Complete document processing pipeline: extract text, chunk, and generate embeddings.
    
    Args:
        file_path: Path to the document file
        document_id: Unique identifier for the document
        chunk_size: Size of text chunks
        overlap: Overlap between chunks
        
    Returns:
        List of processed chunks with embeddings
        
    Raises:
        Exception: If processing fails
    """
    import os
    
    try:
        # Determine file type
        file_extension = os.path.splitext(file_path)[1].lower().strip('.')
        
        # Extract text
        logger.info(f"Extracting text from {file_path}")
        text_content = extract_text_from_file(file_path, file_extension)
        
        if not text_content or not text_content.strip():
            logger.warning(f"No text content extracted from {file_path}")
            return []
        
        # Chunk text
        logger.info(f"Chunking text into {chunk_size} character chunks with {overlap} overlap")
        chunks = chunk_text(text_content, chunk_size=chunk_size, overlap=overlap)
        
        if not chunks:
            logger.warning(f"No chunks created from {file_path}")
            return []
        
        # Extract text content for embedding
        chunk_texts = [chunk['content'] for chunk in chunks]
        
        # Generate embeddings
        logger.info(f"Generating embeddings for {len(chunk_texts)} chunks")
        embeddings = generate_embedding_batch(chunk_texts)
        
        # Combine chunks with embeddings
        processed_chunks = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            processed_chunk = {
                'chunk_id': f"{document_id}_chunk_{i:04d}",
                'document_id': document_id,
                'content': chunk['content'],
                'embedding': embedding,
                'start_position': chunk['start_position'],
                'end_position': chunk['end_position'],
                'chunk_index': i,
                'word_count': chunk.get('word_count', len(chunk['content'].split())),
                'char_count': chunk.get('char_count', len(chunk['content'])),
                'embedding_model': 'amazon.titan-embed-text-v1',
                'file_path': file_path,
                'file_type': file_extension
            }
            processed_chunks.append(processed_chunk)
        
        logger.info(f"Successfully processed {len(processed_chunks)} chunks from {file_path}")
        return processed_chunks
        
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {str(e)}")
        raise


def process_text_for_embedding(text: str, document_id: str, 
                             chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
    """
    Process raw text for embedding: chunk and generate embeddings.
    
    Args:
        text: Raw text content
        document_id: Unique identifier for the document
        chunk_size: Size of text chunks
        overlap: Overlap between chunks
        
    Returns:
        List of processed chunks with embeddings
    """
    try:
        if not text or not text.strip():
            logger.warning("No text content provided")
            return []
        
        # Chunk text
        logger.info(f"Chunking text into {chunk_size} character chunks with {overlap} overlap")
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
        
        if not chunks:
            logger.warning("No chunks created from text")
            return []
        
        # Extract text content for embedding
        chunk_texts = [chunk['content'] for chunk in chunks]
        
        # Generate embeddings
        logger.info(f"Generating embeddings for {len(chunk_texts)} chunks")
        embeddings = generate_embedding_batch(chunk_texts)
        
        # Combine chunks with embeddings
        processed_chunks = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            processed_chunk = {
                'chunk_id': f"{document_id}_chunk_{i:04d}",
                'document_id': document_id,
                'content': chunk['content'],
                'embedding': embedding,
                'start_position': chunk['start_position'],
                'end_position': chunk['end_position'],
                'chunk_index': i,
                'word_count': chunk.get('word_count', len(chunk['content'].split())),
                'char_count': chunk.get('char_count', len(chunk['content'])),
                'embedding_model': 'amazon.titan-embed-text-v1'
            }
            processed_chunks.append(processed_chunk)
        
        logger.info(f"Successfully processed {len(processed_chunks)} chunks")
        return processed_chunks
        
    except Exception as e:
        logger.error(f"Error processing text for embedding: {str(e)}")
        raise


def validate_embedding_vector(embedding: List[float], expected_dimension: int = 1536) -> bool:
    """
    Validate an embedding vector.
    
    Args:
        embedding: Embedding vector to validate
        expected_dimension: Expected dimension of the embedding
        
    Returns:
        True if valid, False otherwise
    """
    if not isinstance(embedding, list):
        return False
    
    if len(embedding) != expected_dimension:
        return False
    
    if not all(isinstance(x, (int, float)) for x in embedding):
        return False
    
    # Check for NaN or infinite values
    import math
    if any(math.isnan(x) or math.isinf(x) for x in embedding):
        return False
    
    return True


def calculate_embedding_similarity(embedding1: List[float], embedding2: List[float]) -> float:
    """
    Calculate cosine similarity between two embedding vectors.
    
    Args:
        embedding1: First embedding vector
        embedding2: Second embedding vector
        
    Returns:
        Cosine similarity score between -1 and 1
    """
    import math
    
    if len(embedding1) != len(embedding2):
        raise ValueError("Embeddings must have the same dimension")
    
    # Calculate dot product
    dot_product = sum(a * b for a, b in zip(embedding1, embedding2))
    
    # Calculate magnitudes
    magnitude1 = math.sqrt(sum(a * a for a in embedding1))
    magnitude2 = math.sqrt(sum(b * b for b in embedding2))
    
    # Avoid division by zero
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
    
    # Calculate cosine similarity
    similarity = dot_product / (magnitude1 * magnitude2)
    
    # Clamp to [-1, 1] to handle floating point errors
    return max(-1.0, min(1.0, similarity))

# OpenSearch utilities for document search and retrieval
def create_opensearch_client(endpoint: str, region: str = 'us-east-1'):
    """
    Create an OpenSearch client for AWS OpenSearch Serverless.
    
    Args:
        endpoint: OpenSearch endpoint URL
        region: AWS region
        
    Returns:
        Configured OpenSearch client
    """
    try:
        from opensearchpy import OpenSearch, RequestsHttpConnection
        from aws_requests_auth.aws_auth import AWSRequestsAuth
        import boto3
        
        credentials = boto3.Session().get_credentials()
        awsauth = AWSRequestsAuth(credentials, region, 'aoss')
        
        client = OpenSearch(
            hosts=[{'host': endpoint.replace('https://', ''), 'port': 443}],
            http_auth=awsauth,
            use_ssl=True,
            verify_certs=True,
            connection_class=RequestsHttpConnection,
            timeout=60
        )
        
        return client
        
    except ImportError as e:
        raise ImportError("opensearch-py and aws-requests-auth packages are required") from e

def search_knowledge_base(opensearch_client, 
                         index_name: str,
                         query_text: str = None,
                         query_embedding: List[float] = None,
                         size: int = 10,
                         min_score: float = 0.7,
                         filters: Dict[str, Any] = None,
                         search_type: str = 'vector') -> Dict[str, Any]:
    """
    Search the knowledge base using various search strategies.
    
    Args:
        opensearch_client: OpenSearch client instance
        index_name: Name of the search index
        query_text: Text query for search
        query_embedding: Pre-computed query embedding
        size: Number of results to return
        min_score: Minimum relevance score threshold
        filters: Additional filters to apply
        search_type: Type of search ('vector', 'keyword', 'hybrid')
        
    Returns:
        Search results with metadata
    """
    try:
        search_body = {
            "size": size,
            "_source": {
                "excludes": ["embedding"]  # Don't return embeddings in results
            }
        }
        
        if search_type == 'vector':
            if not query_embedding:
                raise ValueError("query_embedding is required for vector search")
            
            search_body["query"] = {
                "knn": {
                    "embedding": {
                        "vector": query_embedding,
                        "k": size
                    }
                }
            }
            
        elif search_type == 'keyword':
            if not query_text:
                raise ValueError("query_text is required for keyword search")
            
            search_body["query"] = {
                "multi_match": {
                    "query": query_text,
                    "fields": ["title^2", "chunk_content", "content"],
                    "type": "best_fields"
                }
            }
            
        elif search_type == 'hybrid':
            if not query_embedding or not query_text:
                raise ValueError("Both query_text and query_embedding are required for hybrid search")
            
            search_body["query"] = {
                "bool": {
                    "should": [
                        {
                            "knn": {
                                "embedding": {
                                    "vector": query_embedding,
                                    "k": size,
                                    "boost": 1.0
                                }
                            }
                        },
                        {
                            "multi_match": {
                                "query": query_text,
                                "fields": ["title^2", "chunk_content", "content"],
                                "type": "best_fields",
                                "boost": 0.5
                            }
                        }
                    ]
                }
            }
            
        else:
            raise ValueError(f"Unknown search_type: {search_type}")
        
        # Add filters if provided
        if filters:
            if "query" in search_body:
                # Wrap existing query in bool query with filters
                existing_query = search_body["query"]
                search_body["query"] = {
                    "bool": {
                        "must": [existing_query],
                        "filter": _build_opensearch_filters(filters)
                    }
                }
            else:
                search_body["query"] = {
                    "bool": {
                        "filter": _build_opensearch_filters(filters)
                    }
                }
        
        # Add minimum score threshold
        if min_score > 0:
            search_body["min_score"] = min_score
        
        # Execute search
        response = opensearch_client.search(
            index=index_name,
            body=search_body
        )
        
        # Process results
        hits = response.get("hits", {}).get("hits", [])
        processed_results = []
        
        for hit in hits:
            source = hit["_source"]
            processed_results.append({
                "chunk_id": hit["_id"],
                "score": hit["_score"],
                "document_id": source["document_id"],
                "title": source["title"],
                "authors": source["authors"],
                "chunk_content": source["chunk_content"],
                "start_position": source["start_position"],
                "end_position": source["end_position"],
                "metadata": source.get("metadata", {}),
                "publication_date": source.get("publication_date")
            })
        
        return {
            "total_hits": response.get("hits", {}).get("total", {}).get("value", 0),
            "max_score": response.get("hits", {}).get("max_score", 0),
            "search_type": search_type,
            "results": processed_results
        }
        
    except Exception as e:
        logger.error(f"Error searching knowledge base: {str(e)}")
        return {
            "error": str(e),
            "results": []
        }

def _build_opensearch_filters(filters: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Build OpenSearch filter clauses from filter dictionary"""
    filter_clauses = []
    
    for field, value in filters.items():
        if isinstance(value, list):
            filter_clauses.append({"terms": {field: value}})
        elif isinstance(value, dict):
            if "range" in value:
                filter_clauses.append({"range": {field: value["range"]}})
            elif "exists" in value:
                if value["exists"]:
                    filter_clauses.append({"exists": {"field": field}})
                else:
                    filter_clauses.append({"bool": {"must_not": {"exists": {"field": field}}}})
        else:
            filter_clauses.append({"term": {field: value}})
    
    return filter_clauses

def get_document_by_id(opensearch_client, index_name: str, document_id: str) -> Dict[str, Any]:
    """
    Retrieve all chunks for a specific document.
    
    Args:
        opensearch_client: OpenSearch client instance
        index_name: Name of the search index
        document_id: ID of the document to retrieve
        
    Returns:
        Document data with all chunks
    """
    try:
        search_body = {
            "query": {
                "term": {"document_id": document_id}
            },
            "size": 1000,  # Assuming max 1000 chunks per document
            "sort": [{"start_position": {"order": "asc"}}],
            "_source": {
                "excludes": ["embedding"]
            }
        }
        
        response = opensearch_client.search(
            index=index_name,
            body=search_body
        )
        
        hits = response.get("hits", {}).get("hits", [])
        
        if not hits:
            return {"error": f"Document {document_id} not found"}
        
        # Extract document metadata from first chunk
        first_chunk = hits[0]["_source"]
        
        # Collect all chunks
        chunks = []
        for hit in hits:
            source = hit["_source"]
            chunks.append({
                "chunk_id": hit["_id"],
                "content": source["chunk_content"],
                "start_position": source["start_position"],
                "end_position": source["end_position"]
            })
        
        return {
            "document_id": document_id,
            "title": first_chunk["title"],
            "authors": first_chunk["authors"],
            "publication_date": first_chunk.get("publication_date"),
            "metadata": first_chunk.get("metadata", {}),
            "total_chunks": len(chunks),
            "chunks": chunks
        }
        
    except Exception as e:
        logger.error(f"Error retrieving document {document_id}: {str(e)}")
        return {"error": str(e)}

def delete_document_from_index(opensearch_client, index_name: str, document_id: str) -> Dict[str, Any]:
    """
    Delete all chunks for a document from the index.
    
    Args:
        opensearch_client: OpenSearch client instance
        index_name: Name of the search index
        document_id: ID of the document to delete
        
    Returns:
        Deletion result
    """
    try:
        # First, find all chunks for this document
        search_body = {
            "query": {
                "term": {"document_id": document_id}
            },
            "size": 1000,
            "_source": False  # We only need the IDs
        }
        
        response = opensearch_client.search(
            index=index_name,
            body=search_body
        )
        
        hits = response.get("hits", {}).get("hits", [])
        
        if not hits:
            return {
                "document_id": document_id,
                "status": "not_found",
                "chunks_deleted": 0
            }
        
        # Delete each chunk
        deleted_chunks = []
        for hit in hits:
            chunk_id = hit["_id"]
            try:
                delete_response = opensearch_client.delete(
                    index=index_name,
                    id=chunk_id
                )
                if delete_response["result"] == "deleted":
                    deleted_chunks.append(chunk_id)
            except Exception as e:
                logger.warning(f"Failed to delete chunk {chunk_id}: {str(e)}")
        
        return {
            "document_id": document_id,
            "status": "success",
            "chunks_deleted": len(deleted_chunks),
            "deleted_chunk_ids": deleted_chunks
        }
        
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {str(e)}")
        return {
            "document_id": document_id,
            "status": "error",
            "error": str(e)
        }

def get_index_statistics(opensearch_client, index_name: str) -> Dict[str, Any]:
    """
    Get statistics about the document index.
    
    Args:
        opensearch_client: OpenSearch client instance
        index_name: Name of the search index
        
    Returns:
        Index statistics
    """
    try:
        # Get index stats
        stats_response = opensearch_client.indices.stats(index=index_name)
        
        # Get document count
        count_response = opensearch_client.count(index=index_name)
        
        # Get unique document count
        unique_docs_response = opensearch_client.search(
            index=index_name,
            body={
                "size": 0,
                "aggs": {
                    "unique_documents": {
                        "cardinality": {
                            "field": "document_id"
                        }
                    }
                }
            }
        )
        
        index_stats = stats_response["indices"][index_name]
        
        return {
            "index_name": index_name,
            "total_chunks": count_response["count"],
            "unique_documents": unique_docs_response["aggregations"]["unique_documents"]["value"],
            "index_size_bytes": index_stats["total"]["store"]["size_in_bytes"],
            "index_size_mb": round(index_stats["total"]["store"]["size_in_bytes"] / (1024 * 1024), 2),
            "created_at": index_stats["total"]["indexing"]["index_time_in_millis"]
        }
        
    except Exception as e:
        logger.error(f"Error getting index statistics: {str(e)}")
        return {"error": str(e)}